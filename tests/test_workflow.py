import json
import subprocess
import sys
import tempfile
import unittest
import runpy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader


REPO = Path(__file__).resolve().parents[1]
CLI = REPO / "scripts" / "career_hq.py"


class CareerHQWorkflowTests(unittest.TestCase):
    def run_cli(self, *args, cwd=None, check=True):
        return subprocess.run([sys.executable, str(CLI), *args], cwd=cwd or REPO, text=True, capture_output=True, check=check)

    def test_resume_contact_block_is_wide_centered_and_evenly_distributed(self):
        resume_builder = runpy.run_path(str(REPO / ".agents" / "skills" / "career-hq" / "scripts" / "career_hq.py"))
        evidence = {
            "name": "Avery Rowan",
            "phone": "fictional-phone",
            "location": "Madison, Wisconsin",
            "email": "avery.rowan@example.test",
            "links": [
                {"url": "https://linkedin.example.test/avery-rowan"},
                {"url": "https://github.example.test/avery-rowan"},
                {"url": "https://avery-rowan.example.test"},
            ],
            "summary": "Fictional implementation specialist.",
            "skills": [],
            "experience": [],
            "projects": [],
            "education": [],
        }
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "fictional-resume.docx"
            resume_builder["build_docx"](output, evidence, {"role": "Implementation Specialist"})
            document = Document(output)

        self.assertEqual(len(document.tables), 1)
        self.assertEqual(document.paragraphs[1].text, "Implementation Specialist")
        self.assertEqual(document.paragraphs[1].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        contact_table = document.tables[0]
        self.assertEqual(contact_table.alignment, WD_TABLE_ALIGNMENT.CENTER)
        self.assertFalse(contact_table.autofit)
        self.assertEqual(len(contact_table.rows), 2)
        self.assertTrue(all(len(row.cells) == 3 for row in contact_table.rows))
        self.assertEqual(
            [[cell.text for cell in row.cells] for row in contact_table.rows],
            [
                ["fictional-phone", "Madison, Wisconsin", "avery.rowan@example.test"],
                ["linkedin.example.test/avery-rowan", "github.example.test/avery-rowan", "avery-rowan.example.test"],
            ],
        )
        self.assertTrue(all(round(cell.width.inches, 2) == 2.15 for row in contact_table.rows for cell in row.cells))

        with tempfile.TemporaryDirectory() as folder:
            with self.assertRaisesRegex(SystemExit, "role title is required"):
                resume_builder["build_docx"](Path(folder) / "missing-role.docx", evidence, {})

    def test_sonata_layout_profile_and_education_display_are_consistent(self):
        resume_builder = runpy.run_path(str(REPO / ".agents" / "skills" / "career-hq" / "scripts" / "career_hq.py"))
        profile = {
            "resumePreferences": {
                "layoutProfile": {
                    "value": "sonata-compact-v1",
                    "source": "fictional-fixture",
                    "verifiedAt": "2026-08-26",
                    "verified": True,
                }
            }
        }
        layout = resume_builder["resolve_resume_layout_profile"](profile)
        self.assertEqual(layout["id"], "sonata-compact-v1")
        self.assertEqual(layout["source"], "fictional-fixture")
        self.assertEqual(
            resume_builder["resolve_resume_layout_profile"]({})["id"],
            "baseline-reference-v1",
        )
        self.assertEqual(resume_builder["format_coursework_name"]("TEST 3110 - Systems Design"), "Systems Design")
        self.assertEqual(resume_builder["format_coursework_name"]("Portfolio Studio"), "Portfolio Studio")

        evidence = {
            "name": "Avery Rowan",
            "phone": "fictional-phone",
            "location": "Madison, Wisconsin",
            "email": "avery.rowan@example.test",
            "links": [],
            "summary": "Fictional implementation specialist.",
            "skills": [],
            "experience": [],
            "projects": [],
            "layoutProfile": layout,
            "education": [{
                "degree": "Bachelor of Science",
                "field": "Systems Design",
                "minor": "Operations",
                "institution": "Middle Example State University",
                "start_date": "2020-08",
                "graduation_date": "2024-12",
                "gpa": "3.8",
                "honors": ["Fictional Honors"],
                "coursework": [
                    {"value": "TEST 3110 - Systems Design"},
                    {"value": "Portfolio Studio"},
                ],
            }],
        }
        with tempfile.TemporaryDirectory() as folder:
            docx_path = Path(folder) / "fictional-resume.docx"
            pdf_path = Path(folder) / "fictional-resume.pdf"
            resume_builder["build_docx"](docx_path, evidence, {"role": "Implementation Specialist"})
            resume_builder["build_pdf"](pdf_path, evidence, {"role": "Implementation Specialist"})
            document = Document(docx_path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            education_index = paragraphs.index("EDUCATION")
            self.assertEqual(paragraphs[education_index + 1], "Bachelor of Science in Systems Design, Minor in Operations")
            self.assertEqual(paragraphs[education_index + 2].replace("\u00a0", " "), "Middle Example State University")
            self.assertIn("\u00a0", paragraphs[education_index + 2])
            self.assertEqual(paragraphs[education_index + 3], "August 2020 - December 2024 | GPA: 3.8 | Fictional Honors")
            self.assertIn("Relevant coursework: Systems Design; Portfolio Studio", paragraphs)
            self.assertNotIn("TEST 3110", "\n".join(paragraphs))
            pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(str(pdf_path)).pages)
            self.assertIn("Bachelor of Science in Systems Design, Minor in Operations", pdf_text)
            self.assertIn("Middle Example State University", pdf_text.replace("\u00a0", " "))
            self.assertIn("Relevant coursework: Systems Design; Portfolio Studio", pdf_text)
            self.assertNotIn("TEST 3110", pdf_text)

        metadata = resume_builder["resume_generation_metadata"](layout)
        self.assertEqual(metadata["generator"]["id"], "career-hq.prepare-resume")
        resume_builder["validate_material_format_contract"](metadata)
        with self.assertRaisesRegex(SystemExit, "canonical Career HQ generator"):
            resume_builder["validate_material_format_contract"]({"formatContractVersion": 2, "layoutProfile": layout})

    def test_init_questions_and_conflict_preservation(self):
        with tempfile.TemporaryDirectory() as folder:
            self.run_cli("init", "--workspace", folder)
            questions = json.loads(self.run_cli("questions", "--workspace", folder).stdout)
            self.assertLessEqual(questions["count"], 5)
            self.assertEqual(questions["questions"][0]["field"], "identity.mailingAddress")
            mailing_address = {
                "addressLine1": "100 Fictional Way",
                "city": "Madison",
                "region": "WI",
                "postalCode": "00000",
                "country": "United States",
            }
            self.run_cli(
                "answer",
                "--workspace",
                folder,
                "--field",
                "identity.mailingAddress",
                "--value",
                json.dumps(mailing_address),
                "--source",
                "fictional-test",
            )
            self.run_cli("answer", "--workspace", folder, "--field", "searchDirection.targetRoles", "--value", '["Implementation Specialist"]', "--source", "user-answer:test")
            conflict = json.loads(self.run_cli("answer", "--workspace", folder, "--field", "searchDirection.targetRoles", "--value", '["Program Manager"]', "--source", "resume:test").stdout)
            self.assertFalse(conflict["saved"])
            profile = json.loads((Path(folder) / ".job-search" / "applicant-profile.json").read_text(encoding="utf-8"))
            self.assertEqual(
                profile["identity"]["mailingAddress"]["value"],
                mailing_address,
            )
            self.assertEqual(profile["conflicts"][-1]["status"], "unresolved")

    def test_mailing_address_requires_structured_fields(self):
        with tempfile.TemporaryDirectory() as folder:
            self.run_cli("init", "--workspace", folder)
            result = self.run_cli(
                "answer",
                "--workspace",
                folder,
                "--field",
                "identity.mailingAddress",
                "--value",
                '"100 Fictional Way"',
                "--source",
                "fictional-test",
                check=False,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("must be a JSON object", result.stderr)

    def test_add_job_stores_plain_language_job_summary(self):
        with tempfile.TemporaryDirectory() as folder:
            self.run_cli("init", "--workspace", folder)
            posting = Path(folder) / "fictional-posting.txt"
            posting.write_text(
                "A complete fictional posting for a role that coordinates customer launches.",
                encoding="utf-8",
            )
            summary = "This role coordinates customer launches. You would manage implementation plans and resolve handoff blockers."
            added = json.loads(self.run_cli(
                "add-job",
                "--workspace", folder,
                "--employer", "Fictional Harbor",
                "--role", "Implementation Specialist",
                "--location", "Remote",
                "--arrangement", "Remote",
                "--compensation", "Fictional range",
                "--employment-type", "Full-time",
                "--url", "https://jobs.example.test/fictional-harbor",
                "--posting-file", str(posting),
                "--fit", "strong-match",
                "--job-summary", summary,
                "--strongest-match", "Fictional implementation evidence",
                "--largest-gap", "Fictional gap",
                "--risk", "Fictional risk",
                "--next-action", "Review the role",
                "--current-confirmed",
                "--credible-source",
            ).stdout)
            self.assertEqual(added["jobSummary"], summary)

    def test_submitted_requires_evidence(self):
        with tempfile.TemporaryDirectory() as folder:
            self.run_cli("init", "--workspace", folder)
            ledger_path = Path(folder) / ".job-search" / "applications.json"
            ledger = json.loads(ledger_path.read_text(encoding="utf-8"))
            ledger["applications"].append({
                "id": "test-app", "employer": "Fictional Employer", "role": "Fictional Role",
                "status": "research", "fit": "strong-match", "postingSnapshots": [],
                "materials": [], "nextAction": "Review", "nextActionDate": "2026-07-21",
                "submissionEvidence": None,
            })
            ledger_path.write_text(json.dumps(ledger), encoding="utf-8")
            result = self.run_cli("update-status", "--workspace", folder, "--application-id", "test-app", "--status", "submitted", check=False)
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("requires confirmation evidence", result.stderr)

    def test_review_and_approval_use_employer_and_job_title(self):
        with tempfile.TemporaryDirectory() as folder:
            self.run_cli("init", "--workspace", folder)
            ledger_path = Path(folder) / ".job-search" / "applications.json"
            ledger = json.loads(ledger_path.read_text(encoding="utf-8"))
            ledger["applications"].append({
                "id": "app-random1234",
                "employer": "Fictional Employer",
                "role": "Customer Support Engineer",
                "status": "ready",
                "fit": "strong-match",
                "postingSnapshots": [],
                "materials": [{
                    "version": 1,
                    "files": [{
                        "kind": "pdf",
                        "filename": "app-random1234-resume-v001.pdf",
                        "path": "materials/app-random1234/app-random1234-resume-v001.pdf",
                        "sha256": "fictional",
                    }],
                    "visualVerification": {"status": "passed"},
                }],
                "unresolvedQuestions": [],
                "nextAction": "Review",
                "nextActionDate": "2026-07-21",
                "submissionEvidence": None,
            })
            ledger_path.write_text(json.dumps(ledger), encoding="utf-8")

            review = json.loads(self.run_cli(
                "review",
                "--workspace",
                folder,
                "--application-id",
                "app-random1234",
            ).stdout)
            expected = "I authorize submission to Fictional Employer for Customer Support Engineer"
            self.assertEqual(review["displayName"], "Fictional Employer \u2014 Customer Support Engineer")
            self.assertEqual(review["suggestedApproval"], expected)
            self.assertNotIn("app-random1234", review["suggestedApproval"])
            self.assertIn("Exact wording is not required", review["approvalGuidance"])

            rejected_confirmations = [
                "I authorize submission for app-random1234",
                "Reviewed",
                "Looks good",
                "I approve the resume",
                "Do not submit this application",
            ]
            for confirmation in rejected_confirmations:
                with self.subTest(confirmation=confirmation):
                    rejected = self.run_cli(
                        "approve",
                        "--workspace",
                        folder,
                        "--application-id",
                        "app-random1234",
                        "--confirmation",
                        confirmation,
                        check=False,
                    )
                    self.assertNotEqual(rejected.returncode, 0)
                    self.assertIn("Clear application-specific submission authorization", rejected.stderr)

            approved = json.loads(self.run_cli(
                "approve",
                "--workspace",
                folder,
                "--application-id",
                "app-random1234",
                "--confirmation",
                "Reviewed, authorized for submission",
            ).stdout)
            self.assertEqual(approved["confirmation"], "Reviewed, authorized for submission")
            self.assertEqual(approved["scope"], "app-random1234")

            exact_approval = json.loads(self.run_cli(
                "approve",
                "--workspace",
                folder,
                "--application-id",
                "app-random1234",
                "--confirmation",
                expected,
            ).stdout)
            self.assertEqual(exact_approval["confirmation"], expected)


if __name__ == "__main__":
    unittest.main()
