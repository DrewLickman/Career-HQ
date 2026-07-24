#!/usr/bin/env python3
"""Deterministic private tracker for the Career HQ repository."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import sys
import uuid
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any


FIT_LABELS = {
    "strong-match",
    "reasonable-stretch",
    "low-probability-stretch",
    "not-recommended",
}
STATUSES = {
    "research", "ready", "applied", "assessment", "interview", "offer",
    "rejected", "withdrawn", "closed", "submission-unconfirmed", "submitted",
}
TERMINAL = {"rejected", "withdrawn", "closed"}
QUESTIONS = [
    (
        "identity",
        "identity.mailingAddress",
        "What mailing address should Career HQ use for authorized job applications? "
        "Include street, city, state or province, postal code, and country. You may skip this for now.",
    ),
    ("search-direction", "searchDirection.targetRoles", "Which role titles should Career HQ prioritize?"),
    ("search-direction", "searchDirection.avoidedRoles", "Which roles should it avoid?"),
    ("search-direction", "searchDirection.geography", "What locations and maximum commute work for you?"),
    ("search-direction", "searchDirection.workArrangement", "Which work arrangements are acceptable: remote, hybrid, or on-site?"),
    ("search-direction", "searchDirection.compensation", "What compensation range and employment types are acceptable?"),
    ("career-evidence", "careerEvidence.resumeCorrections", "Do any resume titles or dates need correction?"),
    ("career-evidence", "careerEvidence.approvedMetrics", "Which quantified accomplishments are approved for reuse?"),
    ("career-evidence", "careerEvidence.portfolio", "Which portfolios or projects may support applications?"),
    ("application-defaults", "applicationDefaults.availability", "What availability or start-date language should applications use?"),
    ("application-defaults", "applicationDefaults.salaryStrategy", "What is your preferred salary-answer strategy?"),
    ("application-defaults", "applicationDefaults.references", "What is your reference policy?"),
    ("tracking-preferences", "trackingPreferences.followUpDays", "How many days after submission should the first follow-up occur?"),
    ("tracking-preferences", "trackingPreferences.pipelineStages", "Do you want to customize the default pipeline stages?"),
]


def now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def today_iso() -> str:
    return date.today().isoformat()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_json(path: Path, default: Any = None) -> Any:
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def save_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    temporary.replace(path)


def private_root(workspace: str) -> Path:
    return Path(workspace).resolve() / ".job-search"


def paths_for(workspace: str) -> dict[str, Path]:
    root = private_root(workspace)
    return {
        "root": root,
        "profile": root / "applicant-profile.json",
        "applications": root / "applications.json",
        "postings": root / "postings",
        "materials": root / "materials",
        "reviews": root / "review-packets",
    }


def empty_profile() -> dict[str, Any]:
    return {
        "schemaVersion": 1,
        "createdAt": now_iso(),
        "updatedAt": now_iso(),
        "identity": {},
        "searchDirection": {},
        "careerEvidence": {},
        "applicationDefaults": {},
        "sensitiveAnswers": {},
        "trackingPreferences": {},
        "summary": {},
        "skills": [],
        "experience": [],
        "sources": [],
        "conflicts": [],
    }


def empty_ledger() -> dict[str, Any]:
    return {"schemaVersion": 1, "updatedAt": now_iso(), "applications": []}


def find_resume_sources(workspace: Path) -> list[dict[str, Any]]:
    found: list[dict[str, Any]] = []
    ignored = {".git", ".job-search", "node_modules", "dist", ".vinext", ".next", "sample-data"}
    for path in workspace.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in {".docx", ".pdf"}:
            continue
        if any(part in ignored for part in path.parts):
            continue
        found.append({
            "id": f"resume:{sha256(path)[:12]}",
            "kind": "resume-source",
            "filename": path.name,
            "relativePath": path.relative_to(workspace).as_posix(),
            "sha256": sha256(path),
            "observedAt": now_iso(),
            "needsVerification": True,
        })
    return sorted(found, key=lambda item: item["relativePath"])


def cmd_init(args: argparse.Namespace) -> None:
    workspace = Path(args.workspace).resolve()
    paths = paths_for(args.workspace)
    for key in ("root", "postings", "materials", "reviews"):
        paths[key].mkdir(parents=True, exist_ok=True)
    profile = load_json(paths["profile"], empty_profile())
    known = {source.get("sha256") for source in profile.get("sources", [])}
    discovered = find_resume_sources(workspace)
    profile.setdefault("sources", []).extend(source for source in discovered if source["sha256"] not in known)
    profile["updatedAt"] = now_iso()
    ledger = load_json(paths["applications"], empty_ledger())
    save_json(paths["profile"], profile)
    save_json(paths["applications"], ledger)
    print(json.dumps({
        "workspace": str(paths["root"]),
        "created": [str(path) for path in paths.values()],
        "resumeSourcesFound": len(discovered),
        "next": "Run questions. Career HQ will return at most five unanswered questions.",
    }, indent=2))


def get_nested(payload: dict[str, Any], dotted: str) -> Any:
    current: Any = payload
    for part in dotted.split("."):
        if not isinstance(current, dict) or part not in current:
            return None
        current = current[part]
    return current


def set_nested(payload: dict[str, Any], dotted: str, value: Any) -> None:
    current = payload
    parts = dotted.split(".")
    for part in parts[:-1]:
        current = current.setdefault(part, {})
    current[parts[-1]] = value


def cmd_questions(args: argparse.Namespace) -> None:
    paths = paths_for(args.workspace)
    profile = load_json(paths["profile"])
    if profile is None:
        raise SystemExit("Run init first.")
    unanswered = [
        {"pass": intake_pass, "field": field, "question": prompt}
        for intake_pass, field, prompt in QUESTIONS
        if get_nested(profile, field) is None
    ][:5]
    print(json.dumps({"count": len(unanswered), "questions": unanswered}, indent=2))


def parse_value(raw: str) -> Any:
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return raw


def validate_answer(field: str, proposed: Any) -> None:
    if field != "identity.mailingAddress" or proposed is None:
        return
    required = ("addressLine1", "city", "region", "postalCode", "country")
    if not isinstance(proposed, dict):
        raise SystemExit(
            "identity.mailingAddress must be a JSON object with addressLine1, city, "
            "region, postalCode, and country."
        )
    missing = [
        key for key in required
        if not isinstance(proposed.get(key), str) or not proposed[key].strip()
    ]
    if missing:
        raise SystemExit(
            "identity.mailingAddress is missing required fields: " + ", ".join(missing)
        )


def cmd_answer(args: argparse.Namespace) -> None:
    paths = paths_for(args.workspace)
    profile = load_json(paths["profile"])
    if profile is None:
        raise SystemExit("Run init first.")
    proposed = parse_value(args.value)
    validate_answer(args.field, proposed)
    existing = get_nested(profile, args.field)
    evidence = {
        "value": proposed,
        "source": args.source,
        "verifiedAt": args.verified_at or today_iso(),
        "verified": True,
    }
    if existing is not None and existing.get("value") != proposed:
        conflict = {
            "id": f"conflict-{uuid.uuid4().hex[:10]}",
            "field": args.field,
            "existing": existing,
            "proposed": evidence,
            "detectedAt": now_iso(),
            "status": "resolved-by-explicit-correction" if args.correction else "unresolved",
        }
        profile.setdefault("conflicts", []).append(conflict)
        if not args.correction:
            save_json(paths["profile"], profile)
            print(json.dumps({"saved": False, "conflict": conflict}, indent=2))
            return
    set_nested(profile, args.field, evidence)
    profile["updatedAt"] = now_iso()
    save_json(paths["profile"], profile)
    print(json.dumps({"saved": True, "field": args.field, "evidence": evidence}, indent=2))


def credible_url(url: str) -> bool:
    return bool(re.match(r"^https://[^\s/]+", url))


def cmd_add_job(args: argparse.Namespace) -> None:
    if args.fit not in FIT_LABELS:
        raise SystemExit(f"Invalid fit label: {args.fit}")
    if not args.current_confirmed or not args.credible_source or not credible_url(args.url):
        raise SystemExit("A current HTTPS listing and credible-source confirmation are required.")
    paths = paths_for(args.workspace)
    ledger = load_json(paths["applications"])
    if ledger is None:
        raise SystemExit("Run init first.")
    posting_path = Path(args.posting_file).resolve()
    if not posting_path.is_file():
        raise SystemExit("The complete posting text file is required.")
    posting_bytes = posting_path.read_bytes()
    if not posting_bytes.strip():
        raise SystemExit("The complete posting text file is required.")
    application_id = args.id or f"app-{uuid.uuid4().hex[:10]}"
    if any(item["id"] == application_id for item in ledger["applications"]):
        raise SystemExit(f"Application already exists: {application_id}")
    posting_hash = sha256(posting_path)
    snapshot_name = f"{application_id}-{datetime.now().strftime('%Y%m%dT%H%M%S')}-{posting_hash[:10]}.txt"
    snapshot_path = paths["postings"] / snapshot_name
    snapshot_path.parent.mkdir(parents=True, exist_ok=True)
    with snapshot_path.open("xb") as handle:
        handle.write(posting_bytes)
    next_date = args.next_action_date or today_iso()
    record = {
        "id": application_id,
        "employer": args.employer,
        "role": args.role,
        "location": args.location,
        "arrangement": args.arrangement,
        "compensation": args.compensation,
        "employmentType": args.employment_type,
        "url": args.url,
        "status": "research",
        "fit": args.fit,
        "strongestMatch": args.strongest_match,
        "largestGap": args.largest_gap,
        "risk": args.risk,
        "nextAction": args.next_action,
        "nextActionDate": next_date,
        "postingSnapshots": [{
            "path": snapshot_path.relative_to(paths["root"]).as_posix(),
            "sha256": posting_hash,
            "bytes": snapshot_path.stat().st_size,
            "sourceUrl": args.url,
            "capturedAt": now_iso(),
            "currentConfirmedAt": now_iso(),
            "credibleSourceConfirmed": True,
        }],
        "materials": [], "importantAnswers": [], "unresolvedQuestions": [],
        "approval": None, "submissionEvidence": None,
        "createdAt": now_iso(), "updatedAt": now_iso(),
    }
    ledger["applications"].append(record)
    ledger["updatedAt"] = now_iso()
    save_json(paths["applications"], ledger)
    print(json.dumps(record, indent=2))


def find_application(ledger: dict[str, Any], application_id: str) -> dict[str, Any]:
    for application in ledger.get("applications", []):
        if application.get("id") == application_id:
            return application
    raise SystemExit(f"Unknown application: {application_id}")


def verified_value(item: Any, label: str, required: bool = True) -> Any:
    if not item and not required:
        return None
    if not isinstance(item, dict) or not item.get("verified") or not item.get("source") or not item.get("verifiedAt"):
        raise SystemExit(f"Unverified or untraceable profile value: {label}")
    return item.get("value")


def collect_resume_evidence(profile: dict[str, Any], posting_text: str) -> dict[str, Any]:
    unresolved = [conflict for conflict in profile.get("conflicts", []) if conflict.get("status") == "unresolved"]
    if unresolved:
        raise SystemExit("Resolve profile conflicts before generating materials.")
    name = verified_value(profile.get("identity", {}).get("displayName"), "identity.displayName")
    email = verified_value(profile.get("identity", {}).get("email"), "identity.email", required=False)
    location = verified_value(profile.get("identity", {}).get("location"), "identity.location", required=False)
    summary = verified_value(profile.get("summary"), "summary")
    skills = []
    for index, skill in enumerate(profile.get("skills", [])):
        skills.append({
            "value": verified_value(skill, f"skills[{index}]"),
            "context": skill.get("context", "Verified profile evidence"),
            "source": skill["source"], "verifiedAt": skill["verifiedAt"],
        })
    experiences = []
    for index, experience in enumerate(profile.get("experience", [])):
        claims = []
        for claim_index, claim in enumerate(experience.get("claims", [])):
            claims.append({
                "text": verified_value(claim, f"experience[{index}].claims[{claim_index}]"),
                "source": claim["source"], "verifiedAt": claim["verifiedAt"],
            })
        if claims:
            experiences.append({**{key: experience.get(key) for key in ("id", "employer", "title", "startDate", "endDate", "experienceType")}, "claims": claims})
    if not experiences:
        raise SystemExit("At least one verified experience claim is required.")
    keywords = set(re.findall(r"[a-z]{4,}", posting_text.lower()))
    for experience in experiences:
        experience["relevance"] = sum(word in keywords for word in re.findall(r"[a-z]{4,}", " ".join(claim["text"] for claim in experience["claims"]).lower()))
    experiences.sort(key=lambda item: item["relevance"], reverse=True)
    return {"name": name, "email": email, "location": location, "summary": summary, "skills": skills, "experience": experiences}


def apply_tailoring(evidence: dict[str, Any], tailoring_path: Path | None) -> dict[str, Any]:
    if tailoring_path is None:
        return evidence
    tailoring = load_json(tailoring_path)
    if not tailoring or tailoring.get("truthReviewed") is not True or not tailoring.get("reviewedAt"):
        raise SystemExit("Tailoring must be explicitly truth-reviewed and dated.")
    if tailoring.get("sourceSummary") != evidence["summary"]:
        raise SystemExit("Tailored summary does not reference the exact verified source summary.")
    if tailoring.get("tailoredSummary"):
        evidence["sourceSummary"] = evidence["summary"]
        evidence["summary"] = tailoring["tailoredSummary"]
    source_claims = {claim["text"]: claim for experience in evidence["experience"] for claim in experience["claims"]}
    for replacement in tailoring.get("claims", []):
        source_text = replacement.get("sourceText")
        tailored_text = replacement.get("tailoredText")
        if source_text not in source_claims or not isinstance(tailored_text, str) or not tailored_text.strip():
            raise SystemExit("Every tailored claim must reference one exact verified source claim.")
        source_numbers = set(re.findall(r"\d+(?:\.\d+)?%?", source_text))
        tailored_numbers = set(re.findall(r"\d+(?:\.\d+)?%?", tailored_text))
        if not tailored_numbers.issubset(source_numbers):
            raise SystemExit(f"Tailored claim introduces an unverified number: {tailored_text}")
        source_claims[source_text]["sourceText"] = source_text
        source_claims[source_text]["text"] = tailored_text
        source_claims[source_text]["tailoringReviewedAt"] = tailoring["reviewedAt"]
    requested_skill_order = tailoring.get("skillOrder", [])
    if requested_skill_order:
        skills_by_name = {str(skill["value"]): skill for skill in evidence["skills"]}
        if any(name not in skills_by_name for name in requested_skill_order):
            raise SystemExit("Tailoring requests a skill that is not verified in the profile.")
        ordered = [skills_by_name[name] for name in requested_skill_order]
        ordered.extend(skill for skill in evidence["skills"] if skill["value"] not in requested_skill_order)
        evidence["skills"] = ordered
    evidence["tailoring"] = {"path": str(tailoring_path), "reviewedAt": tailoring["reviewedAt"], "truthReviewed": True}
    return evidence


def next_version(directory: Path, application_id: str) -> int:
    existing = list(directory.glob(f"{application_id}-resume-v*.docx"))
    versions = [int(match.group(1)) for path in existing if (match := re.search(r"-v(\d+)\.docx$", path.name))]
    return max(versions, default=0) + 1


def build_docx(output: Path, evidence: dict[str, Any], application: dict[str, Any]) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise SystemExit("Install requirements.txt before generating resumes.") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial"); normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4); normal.paragraph_format.line_spacing = 1.1
    for style_name, size in (("Heading 1", 13), ("Heading 2", 11)):
        style = styles[style_name]
        style.font.name = "Arial"; style._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial"); style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor(23, 32, 27)
        style.paragraph_format.space_before = Pt(10); style.paragraph_format.space_after = Pt(4)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run(str(evidence["name"])); run.font.name = "Arial"; run._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial"); run.font.size = Pt(22); run.bold = True; run.font.color.rgb = RGBColor(23, 32, 27)
    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER; contact.paragraph_format.space_after = Pt(10)
    contact.add_run(" | ".join(value for value in (evidence.get("email"), evidence.get("location")) if value))

    document.add_heading("Professional summary", level=1)
    document.add_paragraph(str(evidence["summary"]))
    if evidence["skills"]:
        document.add_heading("Relevant skills", level=1)
        skills = document.add_paragraph()
        skills.add_run(" | ".join(str(skill["value"]) for skill in evidence["skills"]))
    document.add_heading("Experience", level=1)
    for experience in evidence["experience"]:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(8); heading.paragraph_format.space_after = Pt(1)
        left = heading.add_run(f"{experience.get('title', '')} | {experience.get('employer', '')}"); left.bold = True
        dates = document.add_paragraph(f"{experience.get('startDate', '')} - {experience.get('endDate', 'Present')} | {experience.get('experienceType', 'professional')}")
        dates.paragraph_format.space_after = Pt(2)
        for claim in experience["claims"]:
            bullet = document.add_paragraph(style="List Bullet")
            bullet.paragraph_format.left_indent = Inches(0.5); bullet.paragraph_format.first_line_indent = Inches(-0.25); bullet.paragraph_format.space_after = Pt(3); bullet.paragraph_format.line_spacing = 1.1
            bullet.add_run(claim["text"])
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"Tailored for {application['employer']} - {application['role']} | Generated {today_iso()}")
    document.save(output)


def build_pdf(output: Path, evidence: dict[str, Any], application: dict[str, Any]) -> None:
    try:
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise SystemExit("Install requirements.txt before generating resumes.") from exc

    styles = getSampleStyleSheet()
    body = ParagraphStyle("ResumeBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=13, spaceAfter=4, textColor=HexColor("#17201B"))
    title = ParagraphStyle("ResumeTitle", parent=body, fontName="Helvetica-Bold", fontSize=22, leading=24, alignment=TA_CENTER, spaceAfter=3)
    contact = ParagraphStyle("ResumeContact", parent=body, fontSize=9.5, alignment=TA_CENTER, spaceAfter=10)
    heading = ParagraphStyle("ResumeHeading", parent=body, fontName="Helvetica-Bold", fontSize=13, leading=15, spaceBefore=9, spaceAfter=4)
    role = ParagraphStyle("ResumeRole", parent=body, fontName="Helvetica-Bold", spaceBefore=6, spaceAfter=1)
    bullet = ParagraphStyle("ResumeBullet", parent=body, leftIndent=18, firstLineIndent=-9, bulletIndent=5, spaceAfter=3)
    story = [Paragraph(str(evidence["name"]), title)]
    story.append(Paragraph(" | ".join(value for value in (evidence.get("email"), evidence.get("location")) if value), contact))
    story.extend([Paragraph("Professional summary", heading), Paragraph(str(evidence["summary"]), body)])
    if evidence["skills"]:
        story.extend([Paragraph("Relevant skills", heading), Paragraph(" | ".join(str(skill["value"]) for skill in evidence["skills"]), body)])
    story.append(Paragraph("Experience", heading))
    for experience in evidence["experience"]:
        story.append(Paragraph(f"{experience.get('title', '')} | {experience.get('employer', '')}", role))
        story.append(Paragraph(f"{experience.get('startDate', '')} - {experience.get('endDate', 'Present')} | {experience.get('experienceType', 'professional')}", body))
        for claim in experience["claims"]:
            story.append(Paragraph(claim["text"], bullet, bulletText="-"))

    document = SimpleDocTemplate(str(output), pagesize=LETTER, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch, title=f"{evidence['name']} - {application['role']}", author="Career HQ")
    document.build(story)


def cmd_prepare_resume(args: argparse.Namespace) -> None:
    paths = paths_for(args.workspace)
    profile = load_json(paths["profile"])
    ledger = load_json(paths["applications"])
    if profile is None or ledger is None:
        raise SystemExit("Run init first.")
    application = find_application(ledger, args.application_id)
    if not application.get("postingSnapshots"):
        raise SystemExit("A complete immutable posting snapshot is required.")
    snapshot = paths["root"] / application["postingSnapshots"][-1]["path"]
    if not snapshot.exists() or sha256(snapshot) != application["postingSnapshots"][-1]["sha256"]:
        raise SystemExit("Posting snapshot is missing or changed.")
    evidence = collect_resume_evidence(profile, snapshot.read_text(encoding="utf-8"))
    evidence = apply_tailoring(evidence, Path(args.tailoring_file).resolve() if args.tailoring_file else None)
    application_dir = paths["materials"] / application["id"]
    application_dir.mkdir(parents=True, exist_ok=True)
    version = next_version(application_dir, application["id"])
    stem = f"{application['id']}-resume-v{version:03d}"
    docx_path = application_dir / f"{stem}.docx"
    pdf_path = application_dir / f"{stem}.pdf"
    build_docx(docx_path, evidence, application)
    build_pdf(pdf_path, evidence, application)
    manifest = {
        "applicationId": application["id"], "version": version,
        "postingSnapshot": application["postingSnapshots"][-1],
        "generatedAt": now_iso(),
        "claims": [{"text": claim["text"], "sourceText": claim.get("sourceText", claim["text"]), "source": claim["source"], "verifiedAt": claim["verifiedAt"], "tailoringReviewedAt": claim.get("tailoringReviewedAt")} for exp in evidence["experience"] for claim in exp["claims"]],
        "skills": evidence["skills"],
        "tailoring": evidence.get("tailoring"),
        "files": [
            {"kind": "docx", "filename": docx_path.name, "path": docx_path.relative_to(paths["root"]).as_posix(), "sha256": sha256(docx_path)},
            {"kind": "pdf", "filename": pdf_path.name, "path": pdf_path.relative_to(paths["root"]).as_posix(), "sha256": sha256(pdf_path)},
        ],
        "visualVerification": {"status": "required", "verifiedAt": None, "notes": None},
    }
    manifest_path = application_dir / f"{stem}-evidence.json"
    save_json(manifest_path, manifest)
    application["materials"].append({**manifest, "manifestPath": manifest_path.relative_to(paths["root"]).as_posix()})
    if application.get("status") not in TERMINAL | {"submitted"}:
        application["status"] = "ready"
        application["nextAction"] = "Render and visually verify DOCX and PDF, then review the application packet."
    application["updatedAt"] = now_iso()
    ledger["updatedAt"] = now_iso(); save_json(paths["applications"], ledger)
    print(json.dumps({"docx": str(docx_path), "pdf": str(pdf_path), "manifest": str(manifest_path), "next": application["nextAction"]}, indent=2))


def cmd_mark_visual(args: argparse.Namespace) -> None:
    paths = paths_for(args.workspace); ledger = load_json(paths["applications"]); application = find_application(ledger, args.application_id)
    if not application.get("materials"):
        raise SystemExit("No material version exists.")
    material = application["materials"][-1]
    for file in material["files"]:
        path = paths["root"] / file["path"]
        if not path.exists() or sha256(path) != file["sha256"]:
            raise SystemExit(f"Material is missing or changed: {file['filename']}")
    material["visualVerification"] = {"status": "passed", "verifiedAt": now_iso(), "notes": args.notes}
    manifest_path = paths["root"] / material["manifestPath"]
    manifest = load_json(manifest_path); manifest["visualVerification"] = material["visualVerification"]; save_json(manifest_path, manifest)
    if application.get("status") not in TERMINAL | {"submitted"}:
        application["nextAction"] = "Create and review the application packet."
    save_json(paths["applications"], ledger)
    print(json.dumps(material["visualVerification"], indent=2))


def cmd_review(args: argparse.Namespace) -> None:
    paths = paths_for(args.workspace); ledger = load_json(paths["applications"]); application = find_application(ledger, args.application_id)
    if not application.get("materials"):
        raise SystemExit("Generate materials first.")
    latest = application["materials"][-1]
    if latest.get("visualVerification", {}).get("status") != "passed":
        raise SystemExit("Visual verification is required before review.")
    packet = {
        "applicationId": application["id"], "employer": application["employer"], "role": application["role"],
        "materials": latest["files"], "importantAnswers": application.get("importantAnswers", []),
        "compensation": application.get("compensation"),
        "workAuthorization": next((item for item in application.get("importantAnswers", []) if item.get("kind") == "work-authorization"), None),
        "unresolvedQuestions": application.get("unresolvedQuestions", []),
        "createdAt": now_iso(), "submissionAuthorized": False,
    }
    packet_path = paths["reviews"] / f"{application['id']}-review-{datetime.now().strftime('%Y%m%dT%H%M%S')}.json"
    save_json(packet_path, packet); application["reviewPacket"] = packet_path.relative_to(paths["root"]).as_posix(); application["nextAction"] = "Resolve questions and obtain application-specific submission authorization."; save_json(paths["applications"], ledger)
    confirmation = f"I authorize submission to {application['employer']} for {application['role']}"
    print(json.dumps({
        "packet": str(packet_path),
        "review": packet,
        "displayName": f"{application['employer']} \u2014 {application['role']}",
        "requiredApproval": confirmation,
    }, indent=2))


def cmd_approve(args: argparse.Namespace) -> None:
    paths = paths_for(args.workspace); ledger = load_json(paths["applications"]); application = find_application(ledger, args.application_id)
    expected = f"I authorize submission to {application['employer']} for {application['role']}"
    if args.confirmation != expected:
        raise SystemExit(f"Exact application-specific confirmation required: {expected}")
    if not application.get("reviewPacket") or application.get("unresolvedQuestions"):
        raise SystemExit("A complete review packet with no unresolved questions is required.")
    application["approval"] = {"confirmation": args.confirmation, "authorizedAt": now_iso(), "scope": application["id"]}
    application["nextAction"] = "User may submit this specific application; capture confirmation evidence afterward."
    save_json(paths["applications"], ledger); print(json.dumps(application["approval"], indent=2))


def cmd_record_submission(args: argparse.Namespace) -> None:
    paths = paths_for(args.workspace); ledger = load_json(paths["applications"]); application = find_application(ledger, args.application_id)
    if not application.get("approval"):
        raise SystemExit("Application-specific authorization is required before recording an attempt.")
    if args.evidence_file:
        source = Path(args.evidence_file).resolve()
        if not source.is_file() or source.stat().st_size == 0:
            raise SystemExit("Confirmation evidence file is missing or empty.")
        evidence_dir = paths["reviews"] / "submission-evidence"; evidence_dir.mkdir(parents=True, exist_ok=True)
        destination = evidence_dir / f"{application['id']}-{datetime.now().strftime('%Y%m%dT%H%M%S')}{source.suffix.lower()}"
        shutil.copy2(source, destination)
        application["submissionEvidence"] = {"path": destination.relative_to(paths["root"]).as_posix(), "sha256": sha256(destination), "recordedAt": now_iso()}
        application["status"] = "submitted"
        follow_days = int(args.follow_up_days or 7)
        application["nextActionDate"] = (date.today() + timedelta(days=follow_days)).isoformat()
        application["nextAction"] = "Follow up on the confirmed submission."
    else:
        application["status"] = "submission-unconfirmed"
        application["submissionEvidence"] = None
        application["nextActionDate"] = today_iso(); application["nextAction"] = "Find confirmation evidence or verify whether submission completed."
    application["updatedAt"] = now_iso(); save_json(paths["applications"], ledger)
    print(json.dumps({"status": application["status"], "submissionEvidence": application["submissionEvidence"], "nextAction": application["nextAction"]}, indent=2))


def cmd_update_status(args: argparse.Namespace) -> None:
    if args.status not in STATUSES:
        raise SystemExit(f"Invalid status: {args.status}")
    paths = paths_for(args.workspace); ledger = load_json(paths["applications"]); application = find_application(ledger, args.application_id)
    if args.status == "submitted" and not application.get("submissionEvidence"):
        raise SystemExit("Submitted requires confirmation evidence.")
    application["status"] = args.status; application["updatedAt"] = now_iso()
    if args.status in TERMINAL:
        application["nextAction"] = "No active action - terminal status."; application["nextActionDate"] = None
    elif args.next_action:
        application["nextAction"] = args.next_action; application["nextActionDate"] = args.next_action_date or today_iso()
    save_json(paths["applications"], ledger); print(json.dumps(application, indent=2))


def cmd_seed_fixture(args: argparse.Namespace) -> None:
    if args.confirmation != "I understand this is fictional demo data":
        raise SystemExit("Exact confirmation required: I understand this is fictional demo data")
    workspace = Path(args.workspace).resolve()
    paths = paths_for(args.workspace)
    if not paths["root"].exists():
        raise SystemExit("Run init first.")
    fixture_path = workspace / "sample-data" / "applicant-profile.json"
    fixture = load_json(fixture_path)
    if not fixture or fixture.get("fixture") is not True:
        raise SystemExit("A clearly marked fictional applicant fixture is required.")
    save_json(paths["profile"], fixture)
    print(json.dumps({"seeded": str(paths["profile"]), "fixture": True}, indent=2))


def cmd_verify(args: argparse.Namespace) -> None:
    paths = paths_for(args.workspace); errors: list[str] = []
    if not paths["root"].exists():
        print(json.dumps({"ok": True, "initialized": False, "note": "Clean clone has no private workspace yet."}, indent=2)); return
    profile = load_json(paths["profile"]); ledger = load_json(paths["applications"])
    if profile is None: errors.append("Missing applicant-profile.json")
    if ledger is None: errors.append("Missing applications.json")
    for application in (ledger or {}).get("applications", []):
        for snapshot in application.get("postingSnapshots", []):
            path = paths["root"] / snapshot["path"]
            if not path.exists() or sha256(path) != snapshot["sha256"]: errors.append(f"Invalid posting snapshot for {application['id']}")
        if application.get("status") == "submitted" and not application.get("submissionEvidence"): errors.append(f"Submitted without evidence: {application['id']}")
        if application.get("status") in TERMINAL and application.get("nextActionDate"): errors.append(f"Terminal application has follow-up: {application['id']}")
        for material in application.get("materials", []):
            for file in material.get("files", []):
                path = paths["root"] / file["path"]
                if not path.exists() or sha256(path) != file["sha256"]: errors.append(f"Invalid material {file.get('filename')} for {application['id']}")
    print(json.dumps({"ok": not errors, "errors": errors, "applications": len((ledger or {}).get("applications", []))}, indent=2))
    if errors: raise SystemExit(1)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Career HQ private local tracker")
    sub = parser.add_subparsers(dest="command", required=True)
    def workspace(command: argparse.ArgumentParser) -> None: command.add_argument("--workspace", default=".")
    p = sub.add_parser("init"); workspace(p); p.set_defaults(func=cmd_init)
    p = sub.add_parser("questions"); workspace(p); p.set_defaults(func=cmd_questions)
    p = sub.add_parser("answer"); workspace(p); p.add_argument("--field", required=True); p.add_argument("--value", required=True); p.add_argument("--source", required=True); p.add_argument("--verified-at"); p.add_argument("--correction", action="store_true"); p.set_defaults(func=cmd_answer)
    p = sub.add_parser("add-job"); workspace(p)
    for name in ("employer", "role", "location", "arrangement", "compensation", "employment-type", "url", "posting-file", "fit", "strongest-match", "largest-gap", "risk", "next-action"):
        p.add_argument(f"--{name}", required=True)
    p.add_argument("--id"); p.add_argument("--next-action-date"); p.add_argument("--current-confirmed", action="store_true"); p.add_argument("--credible-source", action="store_true"); p.set_defaults(func=cmd_add_job)
    p = sub.add_parser("prepare-resume"); workspace(p); p.add_argument("--application-id", required=True); p.add_argument("--tailoring-file"); p.set_defaults(func=cmd_prepare_resume)
    p = sub.add_parser("mark-visual-verification"); workspace(p); p.add_argument("--application-id", required=True); p.add_argument("--notes", required=True); p.set_defaults(func=cmd_mark_visual)
    p = sub.add_parser("review"); workspace(p); p.add_argument("--application-id", required=True); p.set_defaults(func=cmd_review)
    p = sub.add_parser("approve"); workspace(p); p.add_argument("--application-id", required=True); p.add_argument("--confirmation", required=True); p.set_defaults(func=cmd_approve)
    p = sub.add_parser("record-submission"); workspace(p); p.add_argument("--application-id", required=True); p.add_argument("--evidence-file"); p.add_argument("--follow-up-days", type=int); p.set_defaults(func=cmd_record_submission)
    p = sub.add_parser("update-status"); workspace(p); p.add_argument("--application-id", required=True); p.add_argument("--status", required=True); p.add_argument("--next-action"); p.add_argument("--next-action-date"); p.set_defaults(func=cmd_update_status)
    p = sub.add_parser("seed-fixture"); workspace(p); p.add_argument("--confirmation", required=True); p.set_defaults(func=cmd_seed_fixture)
    p = sub.add_parser("verify"); workspace(p); p.set_defaults(func=cmd_verify)
    return parser


def main() -> None:
    args = build_parser().parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
